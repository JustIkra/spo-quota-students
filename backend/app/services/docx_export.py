"""DOCX export helpers."""
from datetime import datetime
from io import BytesIO
from typing import Iterable
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

from app.schemas import OperatorCredential


def build_credentials_docx(items: Iterable[OperatorCredential]) -> BytesIO:
    """Build a .docx file with a table of SPO / login / password credentials."""
    items = list(items)

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = document.add_heading("Учётные данные операторов СПО", level=1)
    title.alignment = 1

    moscow_now = datetime.now(ZoneInfo("Europe/Moscow"))
    document.add_paragraph(f"Сформировано: {moscow_now.strftime('%d.%m.%Y %H:%M')} (МСК)")
    document.add_paragraph(f"Всего записей: {len(items)}")

    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "№"
    header_cells[1].text = "Учреждение"
    header_cells[2].text = "Логин"
    header_cells[3].text = "Пароль"
    for cell in header_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for index, item in enumerate(items, start=1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = item.spo_name
        row[2].text = item.login
        row[3].text = item.password
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
